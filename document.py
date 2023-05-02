'''
Created By: Caleb Malcarne
Program: Invoice Filler 


'''


from docx import Document
from docx.shared import Pt
from inflection import *
from num_to_word import *
from globals import *
import os


class doc():
    
    def __init__(self, data, allow_data, ImagePaths ,outpath):
        self.data = data
        self.allow_data = allow_data
        self.outpath = outpath
        self.ImagePaths = ImagePaths
        
        self.contract_info = contract_info
        self.allowances = allowances
        
    def fillData(self):
        global contract_info
        index = 0
        for item in contract_info:
            if self.data[index] != '':
                contract_info[item] = self.data[index]
            index += 1
            
        print(contract_info)
        
    def formatNum(self, num):
        nums = num.replace('.', '')
        truthLis = [nums[x].isnumeric() for x in range(len(nums))]
        
        if False not in truthLis:
            return "${:,.2f}".format(float(num))
        else:
            return "$0.00" 
        
    def fillAllow(self):
        self.allowances.clear()
        for allow in self.allow_data:
            self.allowances[allow[0]] = allow[1]
     
    def getAllow(self):
        self.fillAllow()
        allows = "Allowances\n"
        totalPrice = 0
        
        for allow in allowances:
            if allowances[allow].isnumeric():
                totalPrice += float(allowances[allow])
            price = self.formatNum(allowances[allow])
            allows += '{:15s} {:15s}\n'.format(str(price), allow)
        allows += '__________________________________________\n'
        if totalPrice <= 0:
            allows += '$0 Grant total for allowances\n'
        else:
            allows += '{:15} Grant total for allowances\n'.format(self.formatNum(str(totalPrice)))

        return allows
    
    def fill_contract(self):
        self.fillData()
        
        document = Document("docs/Contract.docx")
        style = document.styles['Normal']
        font = style.font
        font.name = 'Stymie Lt BT'
        font.size = Pt(11)
        
        date = self.contract_info["date"]
        client = self.contract_info["client"]
        est_num = self.contract_info["estimate_num"]
        start = self.contract_info["start_date"]
        end = self.contract_info["end_date"]
        payment_amount = num_to_words(self.contract_info["estiamte_amount"])
        hourly_rate = self.contract_info["hourly_rate"]
        time = "Is time of the essence in the completion of this contract?  _____Yes  __X___No"


        if int(self.contract_info["time_important"]) > 0:
            time = "Is time of the essence in the completion of this contract?  __X___Yes  _____No"
        deposit_num = self.contract_info["deposit"]
        if deposit_num == '':
            deposit = (depost_percent/100) * (int(self.contract_info["estiamte_amount"]))
            print((depost_percent/100))
            
        elif str(deposit_num).isnumeric() == True:
            deposit = "${:,.2f}".format(float(deposit_num))  
        else:
            deposit = "INVALID INPUT"   
        
        allow_format = self.getAllow()  
        #----------------------------------------------------------------------#  
        document.paragraphs[2].clear()
        document.paragraphs[2].add_run(text = date, style = None)
        
        document.paragraphs[4].clear()
        document.paragraphs[4].add_run(text = client, style = None)
        #----------------------------------------------------------------------#
        est_num_text = f"work outlined in Estimate #{est_num} detailed below in this contract."
        document.paragraphs[7].clear()
        document.paragraphs[7].add_run(text = est_num_text, style = None)    
        #----------------------------------------------------------------------#
        start_date_text = f"Anticipated start date: {start}"
        end_date_text = f"Anticipated completion date: {end}"
        document.paragraphs[17].clear()
        document.paragraphs[17].add_run(text = start_date_text, style = None)
        
        document.paragraphs[18].clear()
        document.paragraphs[18].add_run(text = end_date_text, style = None)
        #----------------------------------------------------------------------#
        document.paragraphs[7].clear()
        document.paragraphs[7].add_run(text = est_num_text, style = None)   
        #----------------------------------------------------------------------#
        document.paragraphs[9].clear()
        document.paragraphs[9].add_run(text = allow_format, style = None)  
        #----------------------------------------------------------------------#
        payment_paragraph_1 = f'''All of the above work is to be completed in a substantial and workmanlike manner for the sum of'''

        payment_paragraph_2 = f' {payment_amount}'
        
        payment_paragraph_3 = f'''
There will be minor alterations to the contract sum based on changes during construction. 
Any alterations or deviation of the work described and as above will be executed upon written 
order for the same and will be added or deducted from the sum quoted in this contract or will be 
performed on a time and material basis at ${hourly_rate}/man/hour.'''

        document.paragraphs[20].clear()
        document.paragraphs[20].add_run(text = payment_paragraph_1, style = None)
        document.paragraphs[20].add_run(text = payment_paragraph_2, style = None).bold = True
        document.paragraphs[20].add_run(text = payment_paragraph_3, style = None)
        
        #----------------------------------------------------------------------#
        document.paragraphs[22].clear()
        document.paragraphs[22].add_run(text = time, style = None)  
        #----------------------------------------------------------------------#
        
        document.paragraphs[25].clear()
        document.paragraphs[25].add_run(text = f"Deposit: {deposit}", style = None).bold = True
        #----------------------------------------------------------------------#
        
        for image in imageLis:
            path = image.getDir()
            if path in self.ImagePaths:
                name = image.getName()
                description = image.getDescription()
                
                paragraph = document.add_paragraph()
                paragraph.add_run(text = name, style = None).bold = True
                
                pic_paragraph = document.add_paragraph()
                run = pic_paragraph.add_run()
                run.add_picture(path)
                
                desc_paragraph = document.add_paragraph()
                desc_paragraph.add_run(text = f"Description: \n{description}\n", style = None)
                
                if("temp" in path):
                    os.remove(path)
            else:
                if("temp" in path):
                    os.remove(path) 
            
    
        document.save(self.outpath + f"/{client} Contract.docx")

