import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QTextEdit, QLineEdit, QLabel, QPushButton, QFileDialog, QVBoxLayout, QHBoxLayout, QWidget
import docx

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Create a widget to hold the labels, fillable fields and buttons
        widget = QWidget(self)
        self.setCentralWidget(widget)

        # Create a layout for the widget
        layout = QVBoxLayout()
        widget.setLayout(layout)

        # Create a list to store the fillable fields
        self.fields = []

        
        
        # Create a label and a fillable field for each paragraph of text that starts with a '*'
        file_path, _ = QFileDialog.getOpenFileName(self, 'Open .docx file', '', 'Word Document (*.docx)')
        if not file_path:
            return
        doc = docx.Document(file_path)  # Replace 'example.docx' with the path to your .docx file


        for p in doc.paragraphs:
            if p.text.startswith('*'):
                # Create a horizontal layout to hold the label and the fillable field
                hlayout = QHBoxLayout()

                # Create a label for the fillable field
                label = QLabel(p.text[1:], self)
                hlayout.addWidget(label)

                # Create a fillable field for the paragraph of text
                field = QLineEdit(self)
                field.setText(p.text[1:])
                hlayout.addWidget(field)
                self.fields.append(field)

                # Add the horizontal layout to the main layout
                layout.addLayout(hlayout)

        # Create a horizontal layout to hold the buttons
        hlayout = QHBoxLayout()

        # Create an "Open" button
        self.open_button = QPushButton('Open', self)
        self.open_button.clicked.connect(self.open_file)
        hlayout.addWidget(self.open_button)

        # Create a "Save" button
        self.save_button = QPushButton('Save', self)
        self.save_button.clicked.connect(self.save_file)
        hlayout.addWidget(self.save_button)

        # Add the horizontal layout to the main layout
        layout.addLayout(hlayout)

    def open_file(self):
        # Open a file dialog to select the .docx file
        file_path, _ = QFileDialog.getOpenFileName(self, 'Open .docx file', '', 'Word Document (*.docx)')
        if not file_path:
            return

        # Read the .docx file and extract the paragraphs of text that start with a '*'
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.startswith('*')]

        # Set the text in the fillable fields
        for i, p in enumerate(paragraphs):
            self.fields[i].setText(p[1:])

    def save_file(self):
        # Open a file dialog to select the .docx file to save
        file_path, _ = QFileDialog.getSaveFileName(self, 'Save .docx file', '', 'Word Document (*.docx)')
        if not file_path:
            return

        # Create a new .docx file
        doc = docx.Document()

        # Add the text from the fillable fields to the .docx file
        for field in self.fields:
            doc.add_paragraph('*' + field.text())

        # Save the .docx file
        doc.save(file_path)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
