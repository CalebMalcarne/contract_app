from docx import Document
from docx.shared import Pt
from inflection import *
from num_to_word import *

#document = Document("Contract.docx")
#paragraph = document.paragraphs[0]
#val = document.tables[0].cell(0, 2).text
#print(document.tables)

contract_info = {
    "date"        : "2/2/2022",
    "client"      : "client",
    "esimate_num" : "A1234",
    "allowances"  : {},
    "start_date"  : "2/2/2022",
    "end_date"    : "2/3/2022",
    "estiamte_amount": 1595.93,
    "hourly_rate" : 125,
    "time_important": False,
    "deposit": 20000,
}

def fill_contract():
    document = Document("docs/Contract.docx")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Stymie Lt BT'
    font.size = Pt(11)
    #document.font.name = 'Stymie Lt BT'
    
    payment_amount = num_to_words(contract_info["estiamte_amount"])
    hourly_rate = contract_info["hourly_rate"]
    payment_paragraph_1 = f'''All of the above work is to be completed in a substantial and workmanlike manner for the sum of'''

    payment_paragraph_2 = f' {payment_amount} '
    
    payment_paragraph_3 = f'''There will be minor alterations to the contract sum based on changes during construction. 
Any alterations or deviation of the work described and as above will be executed upon written 
order for the same and will be added or deducted from the sum quoted in this contract or will be 
performed on a time and material basis at ${hourly_rate}/man/hour.'''

    document.paragraphs[27].clear()
    document.paragraphs[27].add_run(text = payment_paragraph_1, style = None)
    document.paragraphs[27].add_run(text = payment_paragraph_2, style = None).bold = True
    document.paragraphs[27].add_run(text = payment_paragraph_3, style = None)
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    picture = run.add_picture("image.jpg")
        
    document.save("docs/Contract.docx")
    
    

def get_contract_data():

    for i in contract_info:
        if i != "allowances":
            cleaned = humanize(i)
            contract_info[i] = input(f"{cleaned}: ")
            
    def add_allowance():
        allowances = contract_info["allowances"]
        user_allow_type = input("Allowance Type: ")
        user_allow_amount = input("Allowance Amount: ")
        allowances[user_allow_type] = user_allow_amount
        
        next = input("Add another allowance?(Y/N): ")
        if(next == "Y" or next == "y"):
            return 0
        elif(next == "N" or next == "n"):
            return 0
        else:
            add_allowance()    
            
    print("\nAdd Allowances:\n")
    add_allowance()
    print(contract_info["allowances"])
    

#get_contract_data()
fill_contract()

def get_paragraph(num):
    document = Document("docs/Contract.docx")
    paragraph = document.paragraphs[num].text
    return paragraph

def man_select():    
    while(True):
        user_num = int(input("> "))
        print(get_paragraph(user_num))
        print('\n')
        