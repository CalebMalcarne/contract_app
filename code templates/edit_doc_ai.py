import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QTextEdit, QPushButton, QFileDialog
import docx

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Create a text edit widget
        self.text_edit = QTextEdit(self)
        self.text_edit.setGeometry(10, 10, 480, 320)

        # Create an "Open" button
        self.open_button = QPushButton('Open', self)
        self.open_button.setGeometry(10, 340, 100, 30)
        self.open_button.clicked.connect(self.open_file)

        # Create a "Save" button
        self.save_button = QPushButton('Save', self)
        self.save_button.setGeometry(120, 340, 100, 30)
        self.save_button.clicked.connect(self.save_file)

    def open_file(self):
        # Open a file dialog to select the .docx file
        file_path, _ = QFileDialog.getOpenFileName(self, 'Open .docx file', '', 'Word Document (*.docx)')
        if not file_path:
            return

        # Read the .docx file and extract the text
        doc = docx.Document(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs])

        # Set the text in the text edit widget
        self.text_edit.setPlainText(text)

    def save_file(self):
        # Open a file dialog to select the .docx file to save
        file_path, _ = QFileDialog.getSaveFileName(self, 'Save .docx file', '', 'Word Document (*.docx)')
        if not file_path:
            return

        # Create a new .docx file
        doc = docx.Document()

        # Add the text from the text edit widget to the .docx file
        doc.add_paragraph(self.text_edit.toPlainText())

        # Save the .docx file
        doc.save(file_path)


app = QApplication(sys.argv)
window = MainWindow()
window.show()
sys.exit(app.exec_())
